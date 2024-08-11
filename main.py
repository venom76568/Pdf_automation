import json
import os
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert

# 1. Convert PDF to Word
def pdf_to_word(pdf_path, word_output_path):
    cv = Converter(pdf_path)
    cv.convert(word_output_path, start=0, end=None)
    cv.close()

# 2. Fill Word Document with Data from JSON
def fill_word_document(word_path, data):
    doc = Document(word_path)

    # Append data from JSON in front of the existing text
    for key, value in data.items():
        for paragraph in doc.paragraphs:
            if key in paragraph.text:
                # Replace the key with the key followed by its value on the same line
                paragraph.text = paragraph.text.replace(key, f"{key}: {value}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if key in cell.text:
                        # Replace the key with the key followed by its value on the same line
                        cell.text = cell.text.replace(key, f"{key}: {value}")

    doc.save(word_path)

# 3. Convert Word back to PDF
def word_to_pdf(word_path, pdf_output_path):
    convert(word_path, pdf_output_path)

# 4. Integrated Process to Convert, Fill, and Convert Back
def process_pdf_to_filled_pdf(pdf_path, json_data, output_pdf_path):
    word_path = pdf_path.replace(".pdf", ".docx")

    # Convert PDF to Word
    pdf_to_word(pdf_path, word_path)

    # Fill Word document with JSON data
    fill_word_document(word_path, json_data)

    # Convert the filled Word document back to PDF
    word_to_pdf(word_path, output_pdf_path)

# 5. Function to Process Multiple Entries from JSON
def fill_pdfs_from_json(json_path, template_path, output_dir):
    with open(json_path, 'r') as file:
        data_list = json.load(file)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for i, data in enumerate(data_list):
        output_pdf_path = os.path.join(output_dir, f"filled_form_{i + 1}.pdf")
        process_pdf_to_filled_pdf(template_path, data, output_pdf_path)

if __name__ == "__main__":
    json_path = 'data.json'        # Path to your JSON data file
    template_path = 'form2.pdf' # Path to your PDF template
    output_dir = './output_pdfs'   # Directory to save the filled PDFs

    fill_pdfs_from_json(json_path, template_path, output_dir)
    print(f"PDFs generated successfully in {output_dir}")
